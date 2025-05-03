import os
import logging
import fitz  # PyMuPDF
import pandas as pd
import docx

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_excel(file_path):
    """Extract text from an Excel file using pandas."""
    try:
        df = pd.read_excel(file_path)
        # Convert all cells to string and join them
        text = []
        for col in df.columns:
            text.append(str(col))
            for cell in df[col].astype(str):
                if cell != 'nan':
                    text.append(cell)
        return " ".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from Excel: {str(e)}")
        raise Exception(f"Failed to extract text from Excel: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from a Word document using python-docx."""
    try:
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return " ".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        raise Exception(f"Failed to extract text from Word document: {str(e)}")

def extract_text_from_file(file_path, filename):
    """Extract text from various file types based on extension."""
    file_extension = os.path.splitext(filename)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.xlsx', '.xls']:
        return extract_text_from_excel(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")
